import json
from docx import Document
import io
import logging

logger = logging.getLogger(__name__)

def parse_docx(file_buffer):
    """Parse DOCX file from either a file path or a buffer."""
    try:
        if isinstance(file_buffer, str):  # It's a file path
            logger.info(f"Parsing DOCX from path: {file_buffer}")
            doc = Document(file_buffer)
        else:  # It's a buffer
            logger.info(f"Parsing DOCX from buffer, type: {type(file_buffer)}")
            # Ensure the buffer is at the beginning if it's an in-memory file
            if hasattr(file_buffer, 'seek'):
                file_buffer.seek(0)
            doc = Document(io.BytesIO(file_buffer.read()))  # Read buffer content for Document

        # Extract paragraphs
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        logger.info(f"Extracted {len(paragraphs)} paragraphs from DOCX.")

        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        logger.info(f"Extracted {len(tables)} tables from DOCX.")

        # Combine paragraphs and tables into a single text representation
        extracted_text = "\n".join(paragraphs)
        if tables:
            extracted_text += "\n\nTables:\n"
            for table in tables:
                for row in table:
                    extracted_text += "\t".join(row) + "\n"

        logger.info(f"Total extracted text length: {len(extracted_text)} characters.")
        if extracted_text:
            logger.debug(f"Extracted Text (first 500 chars): {extracted_text[:500]}")
        else:
            logger.warning("No text extracted from DOCX.")

        return extracted_text
    except Exception as e:
        logger.error(f"Error parsing DOCX file: {e}", exc_info=True)
        raise Exception(f"Error parsing DOCX file: {str(e)}")


def parse_inputs(paths):
    """Parse input files and return structured data."""
    try:
        # Handle both file paths and file objects
        if isinstance(paths['scenario'], str):
            with open(paths['scenario'], 'r') as f:
                scenario_data = json.load(f)
        else:
            # Ensure the buffer is at the beginning
            if hasattr(paths['scenario'], 'seek'):
                paths['scenario'].seek(0)
            scenario_data = json.loads(paths['scenario'].read().decode('utf-8'))

        # Parse tax returns
        tax_2023_data = parse_docx(paths['2023'])
        tax_2024_data = parse_docx(paths['2024'])

        return {
            "scenario": scenario_data,
            "tax_return_2023": tax_2023_data,
            "tax_return_2024": tax_2024_data
        }
    except Exception as e:
        logger.error(f"Error parsing inputs: {str(e)}", exc_info=True)
        raise Exception(f"Error parsing inputs: {str(e)}")
