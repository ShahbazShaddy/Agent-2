import os
import json
import logging
import datetime
from pathlib import Path
from fpdf import FPDF
from PyPDF2 import PdfReader
import re
from openai import OpenAI
from dotenv import load_dotenv

logger = logging.getLogger(__name__)
load_dotenv()

# Create output directory for reports
REPORTS_DIR = Path("reports")
REPORTS_DIR.mkdir(exist_ok=True)

def parse_previous_tax_return(file_obj):
    """
    Parse a previous year's tax return file (PDF, JSON, or DOCX).
    
    Args:
        file_obj: The uploaded file object
        
    Returns:
        dict: Extracted tax data or error message
    """
    try:
        # Check file extension
        filename = file_obj.name.lower()
        
        if filename.endswith(".json"):
            # Process JSON file
            file_obj.seek(0)  # Reset file pointer
            data = json.loads(file_obj.read().decode('utf-8'))
            return extract_tax_data_from_json(data)
        
        elif filename.endswith(".pdf"):
            # Process PDF file
            file_obj.seek(0)  # Reset file pointer
            return extract_tax_data_from_pdf(file_obj)
        
        elif filename.endswith(".docx"):
            # Process DOCX file
            file_obj.seek(0)  # Reset file pointer
            return extract_tax_data_from_docx(file_obj)
        
        else:
            return {"error": "Unsupported file format. Please upload a PDF, JSON, or DOCX file."}
    
    except Exception as e:
        logger.error(f"Error parsing previous tax return: {str(e)}")
        return {"error": f"Failed to parse tax return: {str(e)}"}

def extract_tax_data_from_json(json_data):
    """
    Extract relevant tax information from JSON data.
    
    Args:
        json_data (dict): The parsed JSON data
        
    Returns:
        dict: Extracted tax data
    """
    try:
        # For now, just return the data as is - we'll use AI to parse it later
        return {
            "source_type": "json",
            "raw_data": json_data
        }
    except Exception as e:
        logger.error(f"Error extracting data from JSON: {str(e)}")
        return {"error": f"Failed to extract data from JSON: {str(e)}"}

def extract_tax_data_from_pdf(file_obj):
    """
    Extract text content from a PDF file.
    
    Args:
        file_obj: The uploaded PDF file object
        
    Returns:
        dict: Extracted text content
    """
    try:
        # Save PDF to temporary file
        temp_path = REPORTS_DIR / f"temp_tax_return_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
        with open(temp_path, "wb") as f:
            f.write(file_obj.read())
            
        # Extract text from PDF
        text_content = ""
        pdf = PdfReader(temp_path)
        
        for page in pdf.pages:
            text_content += page.extract_text() + "\n"
        
        # Remove temp file
        os.remove(temp_path)
        
        return {
            "source_type": "pdf",
            "text_content": text_content,
            "page_count": len(pdf.pages)
        }
    except Exception as e:
        logger.error(f"Error extracting data from PDF: {str(e)}")
        return {"error": f"Failed to extract data from PDF: {str(e)}"}

def extract_tax_data_from_docx(file_obj):
    """
    Extract text content from a DOCX file.
    
    Args:
        file_obj: The uploaded DOCX file object
        
    Returns:
        dict: Extracted text content
    """
    try:
        # Save DOCX to temporary file
        temp_path = REPORTS_DIR / f"temp_tax_return_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        with open(temp_path, "wb") as f:
            f.write(file_obj.read())
            
        # Extract text from DOCX
        from docx import Document
        doc = Document(temp_path)
        text_content = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text_content += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += cell.text + " | "
                text_content += "\n"
        
        # Remove temp file
        os.remove(temp_path)
        
        return {
            "source_type": "docx",
            "text_content": text_content,
            "page_count": len(doc.paragraphs)  # Approximate measure
        }
    except Exception as e:
        logger.error(f"Error extracting data from DOCX: {str(e)}")
        return {"error": f"Failed to extract data from DOCX: {str(e)}"}

def generate_tax_comparison(previous_year_data, current_year_data, client_data):
    """
    Compare previous and current year tax data and generate a PDF report.
    
    Args:
        previous_year_data (dict): Previous year's tax data
        current_year_data (dict): Current year's tax data
        client_data (dict): Client information
        
    Returns:
        dict: Result with PDF report path or error message
    """
    try:
        # Use OpenAI to analyze and compare the tax returns
        comparison_data = analyze_tax_returns_with_ai(
            previous_year_data, 
            current_year_data, 
            client_data
        )
        
        if "error" in comparison_data:
            return comparison_data
        
        # Generate PDF report
        pdf_path = create_comparison_report(comparison_data, client_data)
        
        return {"report_path": pdf_path, "comparison_data": comparison_data}
    
    except Exception as e:
        logger.error(f"Error generating tax comparison: {str(e)}")
        return {"error": f"Failed to generate tax comparison: {str(e)}"}

def analyze_tax_returns_with_ai(previous_year_data, current_year_data, client_data):
    """
    Use OpenAI to analyze and compare tax returns.
    
    Args:
        previous_year_data (dict): Previous year's tax data
        current_year_data (dict): Current year's tax data
        client_data (dict): Client information
        
    Returns:
        dict: AI analysis and structured comparison data
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        logger.error("OpenAI API key not found")
        return {"error": "OpenAI API key not found. Please check your .env file."}
    
    try:
        # Format data based on source type
        if previous_year_data.get("source_type") == "json":
            # For JSON sources, handle potential nested dictionaries
            raw_data = previous_year_data.get("raw_data", {})
            
            # Check for nested dictionaries that might cause issues and flatten them
            flattened_data = {}
            for key, value in raw_data.items():
                if isinstance(value, dict):
                    # For dictionaries, convert to a string representation
                    flattened_data[key] = json.dumps(value)
                else:
                    flattened_data[key] = value
                    
            previous_data_str = json.dumps(flattened_data, indent=2)
        else:
            previous_data_str = previous_year_data.get("text_content", "")
        
        # Initialize OpenAI client
        client = OpenAI(api_key=api_key)
        
        # Create prompt with instructions to return only numeric values
        prompt = f"""
        Extract key tax metrics from the provided previous year tax return and current year tax calculation for a direct comparison.

        CLIENT INFORMATION:
        ```
        {json.dumps(client_data, indent=2)}
        ```

        PREVIOUS YEAR TAX RETURN:
        ```
        {previous_data_str}
        ```

        CURRENT YEAR TAX CALCULATION:
        ```
        {json.dumps(current_year_data, indent=2)}
        ```

        Create a structured JSON with ONLY key tax metrics for a comparison table.
        Focus on extracting numeric values for direct comparison.
        
        IMPORTANT: 
        1. DO NOT include lengthy analysis text
        2. ONLY include the metrics, values, and minimal descriptions
        3. ALL values for previous_year, current_year, and difference MUST be numeric (integers or floats)
        4. If a value is a dictionary or complex object, calculate its sum or use its most relevant numeric value
        5. NEVER return dictionaries, lists, or complex objects in the numeric fields

        Format your response as a JSON with EXACTLY this structure:
        {{
          "year_labels": ["Previous Year", "Current Year"],
          "key_metrics": [
            {{
              "label": "Metric Name",
              "previous_year": [numeric value only],
              "current_year": [numeric value only],
              "difference": [numeric difference]
            }}
          ]
        }}

        Include these key metrics if available:
        1. Total Income
        2. Adjusted Gross Income
        3. Taxable Income
        4. Total Deductions
        5. Federal Tax
        6. State/Local Tax
        7. Total Tax
        8. Effective Tax Rate (as percentage)
        9. Tax Credits
        10. Tax Refund/Amount Due

        REMEMBER: For each metric, ensure all values are numeric:
        - 'previous_year': must be a number
        - 'current_year': must be a number
        - 'difference': must be a number (current_year minus previous_year)

        If you encounter a dictionary like {{'taxes': 10000, 'interest': 2469}}, use the sum (12469) or most relevant value.
        """
        
        # Call OpenAI API
        logger.info("Calling OpenAI for tax return comparison...")
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a tax analysis system that extracts key metrics and creates structured table data. Return ONLY the requested JSON structure with no additional text."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1  # Lower temperature for more consistent structured output
        )
        
        # Extract and parse response
        content = response.choices[0].message.content
        
        # Extract JSON
        json_start = content.find('{')
        json_end = content.rfind('}') + 1
        
        if json_start >= 0 and json_end > json_start:
            json_content = content[json_start:json_end]
            comparison_data = json.loads(json_content)
            
            logger.info("Successfully extracted tax comparison metrics with AI")
            return comparison_data
        else:
            logger.error("Failed to extract valid JSON from OpenAI response")
            return {"error": "Could not extract comparison data from AI response"}
    
    except Exception as e:
        logger.error(f"Error analyzing tax returns with AI: {str(e)}")
        return {"error": f"Failed to analyze tax returns: {str(e)}"}

def create_comparison_report(comparison_data, client_data):
    """
    Create a clean, table-focused PDF report of the tax return comparison.
    
    Args:
        comparison_data (dict): The comparison data from AI analysis
        client_data (dict): Client information
        
    Returns:
        str: Path to the generated PDF report
    """
    try:
        # Generate a PDF report with FPDF
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Add title page
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, "Tax Comparison Report", 0, 1, "C")
        
        # Add client name if available
        client_name = client_data.get("name", "")
        if not client_name and "ClientDetails" in client_data:
            client_name = client_data["ClientDetails"].get("name", "")
        
        if client_name:
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, f"For: {client_name}", 0, 1, "C")
        
        # Add date
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 10, f"Generated on: {datetime.datetime.now().strftime('%B %d, %Y')}", 0, 1, "C")
        
        # Add comparison table - focus on making this clean and readable
        pdf.ln(10)
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "Tax Year Comparison", 0, 1, "C")
        
        # Create table headers
        pdf.set_fill_color(220, 220, 220)
        pdf.set_font("Arial", "B", 10)
        
        # Get year labels
        year_labels = comparison_data.get("year_labels", ["Previous Year", "Current Year"])
        
        # Set column widths for better readability
        metric_width = 55
        year_width = 35
        diff_width = 35
        percent_width = 35
        
        # Draw table header
        pdf.cell(metric_width, 10, "Tax Metric", 1, 0, "L", True)
        pdf.cell(year_width, 10, year_labels[0], 1, 0, "C", True)
        pdf.cell(year_width, 10, year_labels[1], 1, 0, "C", True)
        pdf.cell(diff_width, 10, "Difference", 1, 0, "C", True)
        pdf.cell(percent_width, 10, "% Change", 1, 1, "C", True)
        
        # Add data rows
        pdf.set_font("Arial", "", 10)
        
        # Keep track of alternating row colors for readability
        row_count = 0
        
        for metric in comparison_data.get("key_metrics", []):
            # Add subtle alternating row background for readability
            if row_count % 2 == 0:
                pdf.set_fill_color(245, 245, 245)
                has_fill = True
            else:
                has_fill = False
            row_count += 1
            
            # Extract values
            label = metric.get("label", "")
            prev_year = metric.get("previous_year", 0)
            curr_year = metric.get("current_year", 0)
            diff = metric.get("difference", 0)
            perc = metric.get("percentage_change", 0)
            
            # Render label cell
            pdf.cell(metric_width, 8, label, 1, 0, "L", has_fill)
            
            # Format and render previous year value
            if isinstance(prev_year, (int, float)) and "rate" not in label.lower():
                pdf.cell(year_width, 8, f"${prev_year:,.2f}", 1, 0, "R", has_fill)
            elif isinstance(prev_year, (int, float)) and "rate" in label.lower():
                pdf.cell(year_width, 8, f"{prev_year:.2f}%", 1, 0, "R", has_fill)
            else:
                pdf.cell(year_width, 8, str(prev_year), 1, 0, "R", has_fill)
            
            # Format and render current year value
            if isinstance(curr_year, (int, float)) and "rate" not in label.lower():
                pdf.cell(year_width, 8, f"${curr_year:,.2f}", 1, 0, "R", has_fill)
            elif isinstance(curr_year, (int, float)) and "rate" in label.lower():
                pdf.cell(year_width, 8, f"{curr_year:.2f}%", 1, 0, "R", has_fill)
            else:
                pdf.cell(year_width, 8, str(curr_year), 1, 0, "R", has_fill)
            
            # Format difference with appropriate coloring
            if isinstance(diff, (int, float)):
                # Determine if this is better when higher (income) or lower (taxes)
                is_better_when_higher = not any(tax_word in label.lower() 
                                              for tax_word in ["tax", "liability"])
                
                # Set color based on positive/negative change and metric type
                if (diff > 0 and is_better_when_higher) or (diff < 0 and not is_better_when_higher):
                    pdf.set_text_color(0, 128, 0)  # Green for good change
                elif (diff < 0 and is_better_when_higher) or (diff > 0 and not is_better_when_higher):
                    pdf.set_text_color(255, 0, 0)  # Red for bad change
                
                # Format the difference value depending on if it's a rate or currency
                sign = "+" if diff > 0 else ""
                if "rate" in label.lower():
                    diff_str = f"{sign}{diff:.2f}%"
                else:
                    diff_str = f"{sign}${diff:,.2f}"
                    
                pdf.cell(diff_width, 8, diff_str, 1, 0, "R", has_fill)
            else:
                pdf.cell(diff_width, 8, "N/A", 1, 0, "R", has_fill)
            
            # Format percentage change
            if isinstance(perc, (int, float)):
                sign = "+" if perc > 0 else ""
                pdf.cell(percent_width, 8, f"{sign}{perc:.2f}%", 1, 1, "R", has_fill)
            else:
                pdf.cell(percent_width, 8, "N/A", 1, 1, "R", has_fill)
            
            # Reset text color
            pdf.set_text_color(0, 0, 0)
        
        # Add legend at bottom of table using rectangles instead of Unicode characters
        pdf.ln(5)
        pdf.set_font("Arial", "", 8)
        
        # Draw green rectangle for favorable change
        pdf.set_fill_color(0, 128, 0)  # Green
        pdf.rect(pdf.get_x(), pdf.get_y(), 5, 5, style="F")
        pdf.set_x(pdf.get_x() + 10)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(70, 5, "Favorable Change", 0, 0)
        
        # Draw red rectangle for unfavorable change
        pdf.set_x(pdf.get_x() + 10)
        pdf.set_fill_color(255, 0, 0)  # Red
        pdf.rect(pdf.get_x(), pdf.get_y(), 5, 5, style="F")
        pdf.set_x(pdf.get_x() + 10)
        pdf.cell(70, 5, "Unfavorable Change", 0, 1)
        
        # Reset fill color
        pdf.set_fill_color(220, 220, 220)
        
        # Save the PDF
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_client_name = "".join(c if c.isalnum() else "_" for c in client_name) if client_name else "tax_client"
        pdf_path = str(REPORTS_DIR / f"{safe_client_name}_Tax_Comparison_{timestamp}.pdf")
        
        pdf.output(pdf_path)
        logger.info(f"Generated table-focused tax comparison report: {pdf_path}")
        
        return pdf_path
        
    except Exception as e:
        logger.error(f"Error creating comparison report: {str(e)}")
        raise Exception(f"Failed to create comparison report: {str(e)}")
